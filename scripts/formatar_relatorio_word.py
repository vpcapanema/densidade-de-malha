# -*- coding: utf-8 -*-
"""Formata o Relatório Word - converte listas em tabelas e adiciona gráficos

Este script:
1. Lê o documento Word existente
2. Identifica trechos com "Top 10" e "Bottom 10" 
3. Converte esses trechos em tabelas formatadas
4. Adiciona gráficos de barras verticais
5. Salva o documento formatado

Uso:
  python scripts/formatar_relatorio_word.py
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
INPUT_DOCX = BASE_DIR / "relatorio" / "Relatorio_Densidade_de_Malha.docx"
OUTPUT_DOCX = BASE_DIR / "relatorio" / "Relatorio_Densidade_de_Malha_Formatado.docx"
FIG_DIR = BASE_DIR / "relatorio" / "figuras_word"


def extrair_dados_municipios(texto: str) -> tuple[list[tuple[str, float]], list[tuple[str, float]]]:
    """Extrai dados de Top 10 e Bottom 10 do texto"""
    
    # Padrão melhorado: Nome do município (valor km) ou (valor km/10000km²) ou (valor km/10.000hab)
    # Aceita nomes com acentos, espaços, hífens
    pattern = r'([A-ZÀ-Ý][A-Za-zÀ-ÿ\s\-]+?)\s+\((\d+[\.,]?\d*)\s*km[^\)]*\)'
    
    # Extrair todos os matches
    matches = re.findall(pattern, texto)
    
    result = []
    for nome, valor in matches:
        nome = nome.strip()
        # Converter vírgula para ponto
        valor = float(valor.replace(',', '.'))
        result.append((nome, valor))
    
    # Se o texto contém "Top 10", retorna como top
    if "Top 10" in texto:
        return result, []
    # Se contém "Bottom 10", retorna como bottom
    elif "Bottom 10" in texto:
        return [], result
    else:
        return result, []


def criar_tabela_municipios(doc: Document, dados: list[tuple[str, float]], titulo: str, cor_header: tuple[int, int, int], unidade: str = "km") -> None:
    """Cria uma tabela formatada no documento Word"""
    
    # Adicionar título
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    # Criar tabela (cabeçalho + dados)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Configurar cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Posição'
    hdr_cells[1].text = 'Município'
    hdr_cells[2].text = f'Valor ({unidade})'
    
    # Formatar cabeçalho
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Cor de fundo
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="{:02x}{:02x}{:02x}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            *cor_header))
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Adicionar dados
    for idx, (municipio, valor) in enumerate(dados, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = municipio
        
        # Formatar valor de acordo com a magnitude
        if valor >= 1:
            row_cells[2].text = f"{valor:.1f}"
        else:
            row_cells[2].text = f"{valor:.3f}"
        
        # Alinhar colunas
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Ajustar largura das colunas
    for row in table.rows:
        row.cells[0].width = Cm(2.0)
        row.cells[1].width = Cm(8.0)
        row.cells[2].width = Cm(3.5)
    
    doc.add_paragraph()  # Espaçamento


def criar_grafico_barras(dados: list[tuple[str, float]], titulo: str, cor: str, output_path: Path) -> None:
    """Cria um gráfico de barras verticais"""
    
    if not dados:
        return
    
    municipios = [m for m, v in dados]
    valores = [v for m, v in dados]
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    bars = ax.bar(range(len(municipios)), valores, color=cor, edgecolor='black', linewidth=0.5)
    
    # Adicionar valores sobre as barras
    for i, (bar, valor) in enumerate(zip(bars, valores)):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height + max(valores) * 0.01,
                f'{valor:.1f}',
                ha='center', va='bottom', fontsize=9, fontweight='bold')
    
    # Configurar eixos
    ax.set_xticks(range(len(municipios)))
    ax.set_xticklabels(municipios, rotation=45, ha='right', fontsize=9)
    ax.set_ylabel('Extensão (km)', fontsize=11, fontweight='bold')
    ax.set_title(titulo, fontsize=13, fontweight='bold', pad=20)
    
    # Grid
    ax.yaxis.grid(True, linestyle='--', alpha=0.3)
    ax.set_axisbelow(True)
    
    # Ajustar layout
    plt.tight_layout()
    
    # Salvar
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close(fig)


def processar_top_bottom(doc: Document, idx_top: int, idx_bottom: int) -> bool:
    """Processa um par de parágrafos Top 10 / Bottom 10"""
    
    texto_top = doc.paragraphs[idx_top].text
    texto_bottom = doc.paragraphs[idx_bottom].text
    
    print(f"\n  📍 Processando parágrafos {idx_top} (Top 10) e {idx_bottom} (Bottom 10)")
    
    # Extrair dados
    top_10, _ = extrair_dados_municipios(texto_top)
    _, bottom_10 = extrair_dados_municipios(texto_bottom)
    
    if not top_10 or not bottom_10:
        print(f"  ⚠ Dados incompletos: Top={len(top_10)}, Bottom={len(bottom_10)}")
        return False
    
    print(f"  ✓ Extraídos {len(top_10)} municípios Top e {len(bottom_10)} Bottom")
    
    # Determinar qual indicador baseado no contexto
    indicador_nome = "Extensão Total"
    unidade = "km"
    cor_top = (43, 140, 190)  # Azul
    cor_bottom = (222, 45, 38)  # Vermelho
    
    contexto = texto_top.lower() + texto_bottom.lower()
    
    # Identificar pela unidade no texto
    if "km/10.000hab" in texto_top or "km/10.000hab" in texto_bottom:
        indicador_nome = "Densidade por População"
        unidade = "km/10.000hab"
        cor_top = (230, 85, 13)  # Laranja
    elif "km/10000km²" in texto_top or "km/10000km²" in texto_bottom:
        indicador_nome = "Densidade por Área"
        unidade = "km/10000km²"
        cor_top = (43, 140, 190)  # Azul (mantém)
    
    # Criar gráficos
    fig_top_path = FIG_DIR / f"grafico_top10_{indicador_nome.lower().replace(' ', '_').replace('/', '_')}.png"
    criar_grafico_barras(
        top_10,
        f"Top 10 Municípios - {indicador_nome}",
        '#2b8cbe',
        fig_top_path
    )
    print(f"  ✓ Gráfico Top 10 criado: {fig_top_path.name}")
    
    fig_bottom_path = FIG_DIR / f"grafico_bottom10_{indicador_nome.lower().replace(' ', '_').replace('/', '_')}.png"
    criar_grafico_barras(
        bottom_10,
        f"Bottom 10 Municípios - {indicador_nome}",
        '#de2d26',
        fig_bottom_path
    )
    print(f"  ✓ Gráfico Bottom 10 criado: {fig_bottom_path.name}")
    
    # Limpar os parágrafos originais
    doc.paragraphs[idx_top].text = ""
    doc.paragraphs[idx_bottom].text = ""
    
    # Inserir conteúdo após o parágrafo Bottom (para manter ordem)
    p_element = doc.paragraphs[idx_bottom]._element
    parent = p_element.getparent()
    p_index = parent.index(p_element)
    
    # Criar novo documento temporário para gerar as tabelas
    temp_doc = Document()
    
    # Adicionar tabelas e gráficos
    criar_tabela_municipios(temp_doc, top_10, f"Top 10 Municípios - {indicador_nome}", cor_top, unidade)
    
    # Adicionar gráfico Top 10
    if fig_top_path.exists():
        temp_doc.add_paragraph()
        p_img = temp_doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(str(fig_top_path), width=Cm(15))
    
    temp_doc.add_paragraph()  # Espaçamento entre tabelas
    
    criar_tabela_municipios(temp_doc, bottom_10, f"Bottom 10 Municípios - {indicador_nome}", cor_bottom, unidade)
    
    # Adicionar gráfico Bottom 10
    if fig_bottom_path.exists():
        temp_doc.add_paragraph()
        p_img = temp_doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(str(fig_bottom_path), width=Cm(15))
    
    # Copiar elementos do documento temporário para o documento principal
    for element in temp_doc.element.body:
        # Inserir após o parágrafo Bottom
        parent.insert(p_index + 1, element)
        p_index += 1
    
    return True


def formatar_documento():
    """Função principal para formatar o documento"""
    
    print(f"\n{'='*60}")
    print("FORMATAÇÃO DO RELATÓRIO WORD")
    print(f"{'='*60}\n")
    
    # Verificar se arquivo existe
    if not INPUT_DOCX.exists():
        print(f"❌ Arquivo não encontrado: {INPUT_DOCX}")
        return
    
    print(f"📄 Lendo documento: {INPUT_DOCX.name}")
    
    # Carregar documento
    doc = Document(INPUT_DOCX)
    
    print(f"✓ Documento carregado com {len(doc.paragraphs)} parágrafos")
    print(f"\n🔍 Procurando por pares Top 10 / Bottom 10...\n")
    
    # Encontrar pares de Top 10 / Bottom 10
    pares = []
    i = 0
    while i < len(doc.paragraphs):
        texto = doc.paragraphs[i].text
        if "Top 10 municípios" in texto and "maiores valores" in texto:
            # Procurar Bottom 10 logo após
            for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
                texto_next = doc.paragraphs[j].text
                if "Bottom 10 municípios" in texto_next and "menores valores" in texto_next:
                    pares.append((i, j))
                    print(f"  ✓ Par encontrado: parágrafos {i} e {j}")
                    break
        i += 1
    
    print(f"\n📊 Total de pares encontrados: {len(pares)}")
    
    # Processar pares (de trás para frente para não afetar índices)
    modificados = 0
    for idx_top, idx_bottom in reversed(pares):
        if processar_top_bottom(doc, idx_top, idx_bottom):
            modificados += 1
    
    # Salvar documento formatado
    if modificados > 0:
        print(f"\n💾 Salvando documento formatado: {OUTPUT_DOCX.name}")
        doc.save(OUTPUT_DOCX)
        print(f"✓ Documento salvo com sucesso!")
        print(f"\n📊 Estatísticas:")
        print(f"  - Seções processadas: {modificados}")
        print(f"  - Gráficos gerados: {len(list(FIG_DIR.glob('grafico_*.png')))}")
        print(f"\n✅ Formatação concluída!")
        print(f"\n📂 Arquivos gerados:")
        print(f"  - Documento: {OUTPUT_DOCX}")
        print(f"  - Gráficos: {FIG_DIR}")
    else:
        print("\n⚠ Nenhuma seção Top/Bottom encontrada para processar")
    
    print(f"\n{'='*60}\n")


if __name__ == "__main__":
    formatar_documento()
