# -*- coding: utf-8 -*-
"""Gerador de Relatório Word (DOCX) — Densidade de Malha

Gera um documento Word técnico, com:
- Introdução curta
- Objetivos (a partir das questões norteadoras)
- Dados e metodologia (breve; detalhamento no anexo)
- Resultados:
  - Visão geral do estado
  - Indicadores (equação + gráfico Top 10 / Bottom 10 + mapa + breve discussão)
- Conclusão
- Anexo metodológico

Saída:
- DOCX em relatorio/Relatorio_Densidade_de_Malha.docx
- Figuras em relatorio/figuras_word/

Uso:
  "D:/densidade _de_malha/.venv/Scripts/python.exe" scripts/gerar_relatorio_word.py

Obs.: Este script assume que existe `dados/saida/densidade_malha.gpkg`.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

import geopandas as gpd
import matplotlib.pyplot as plt
import pandas as pd

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


BASE_DIR = Path(__file__).resolve().parents[1]
GPKG_PATH = BASE_DIR / "dados" / "saida" / "densidade_malha.gpkg"
OUT_DOCX = BASE_DIR / "relatorio" / "Relatorio_Densidade_de_Malha.docx"
FIG_DIR = BASE_DIR / "relatorio" / "figuras_word"


@dataclass(frozen=True)
class Indicador:
    key: str
    titulo: str
    coluna_valor: str
    unidade: str
    equacao: str
    descricao: str
    cmap: str


INDICADORES: list[Indicador] = [
    Indicador(
        key="ext_total",
        titulo="Extensão total de rodovias",
        coluna_valor="Ext_Total",
        unidade="km",
        equacao="Extensão total (km) = Σ extensão dos trechos na unidade territorial",
        descricao=(
            "Mede a quantidade absoluta de infraestrutura rodoviária no município. "
            "É um indicador de porte/volume de rede e tende a ser maior em municípios extensos "
            "e/ou com maior número de eixos estruturantes."
        ),
        cmap="Purples",
    ),
    Indicador(
        key="dens_area",
        titulo="Densidade de malha por área",
        coluna_valor="Dens_Area_Total",
        unidade="km/10000km²",
        equacao="Densidade por área (km/10000km²) = (Extensão total (km) / Área (km²)) × 10.000",
        descricao=(
            "Indica a cobertura territorial da infraestrutura. Valores maiores sugerem maior "
            "capilaridade/serviço por área; valores baixos podem sinalizar potenciais gargalos "
            "de cobertura territorial, dependendo do contexto físico e socioeconômico."
        ),
        cmap="Blues",
    ),
    Indicador(
        key="dens_pop",
        titulo="Densidade de malha por população (disponibilidade)",
        coluna_valor="Dens_Pop_Total",
        unidade="km/10.000hab",
        equacao="Densidade por população (km/10.000hab) = Extensão total (km) / (População / 10.000)",
        descricao=(
            "Indica a disponibilidade de infraestrutura em relação ao contingente populacional. "
            "Valores elevados podem ocorrer em municípios pouco populosos (inclusive rurais), "
            "enquanto valores baixos sugerem maior pressão populacional sobre a rede existente."
        ),
        cmap="Oranges",
    ),
]


def _fmt_num(value: float, decimals: int = 2) -> str:
    s = f"{value:,.{decimals}f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt_int(value: float) -> str:
    s = f"{value:,.0f}"
    return s.replace(",", ".")


def carregar_municipios() -> gpd.GeoDataFrame:
    if not GPKG_PATH.exists():
        raise FileNotFoundError(
            f"GeoPackage não encontrado: {GPKG_PATH}. "
            "Gere os dados antes (scripts/02_processar_dados.py)."
        )

    gdf = gpd.read_file(GPKG_PATH, layer="municipios")

    required = {"NM_MUN", "Area_km2", "Populacao"} | {i.coluna_valor for i in INDICADORES}
    missing = sorted([c for c in required if c not in gdf.columns])
    if missing:
        raise ValueError(f"Colunas ausentes na camada 'municipios': {missing}")

    # Garantir tipos
    for col in ["Area_km2", "Populacao"] + [i.coluna_valor for i in INDICADORES]:
        gdf[col] = pd.to_numeric(gdf[col], errors="coerce")

    return gdf


def gerar_grafico_top_bottom(
    gdf: gpd.GeoDataFrame,
    indicador: Indicador,
    out_path: Path,
    label_col: str = "NM_MUN",
    n: int = 10,
) -> None:
    df = gdf[[label_col, indicador.coluna_valor]].dropna().copy()

    # Garantir numérico (evita comportamento estranho de plot em colunas object)
    df[indicador.coluna_valor] = pd.to_numeric(df[indicador.coluna_valor], errors="coerce")
    df = df.dropna(subset=[indicador.coluna_valor])

    df_top = df.nlargest(n, indicador.coluna_valor).sort_values(indicador.coluna_valor)
    df_bot = df.nsmallest(n, indicador.coluna_valor).sort_values(indicador.coluna_valor)

    fig = plt.figure(figsize=(16, 8))
    gs = fig.add_gridspec(1, 2, wspace=0.25)

    overall_max = float(df[indicador.coluna_valor].max()) if len(df) else 0.0

    ax1 = fig.add_subplot(gs[0, 0])
    bars1 = ax1.barh(df_top[label_col], df_top[indicador.coluna_valor], color="#2b8cbe")
    ax1.set_title(f"Top {n} — {indicador.titulo}")
    ax1.set_xlabel(indicador.unidade)

    max_top = float(df_top[indicador.coluna_valor].max()) if len(df_top) else 0.0
    x_max_top = max_top if max_top > 0 else (overall_max * 0.10 if overall_max > 0 else 1.0)
    ax1.set_xlim(0, x_max_top * 1.15)
    offset_top = x_max_top * 0.02
    for rect, v in zip(bars1, df_top[indicador.coluna_valor]):
        ax1.text(float(v) + offset_top, rect.get_y() + rect.get_height() / 2, _fmt_num(float(v), 3),
                 va="center", fontsize=8)

    ax2 = fig.add_subplot(gs[0, 1])
    bars2 = ax2.barh(df_bot[label_col], df_bot[indicador.coluna_valor], color="#de2d26")
    ax2.set_title(f"Bottom {n} — {indicador.titulo}")
    ax2.set_xlabel(indicador.unidade)

    # Caso comum: muitos mínimos iguais a 0, o que pode fazer o gráfico parecer "em branco".
    max_bot = float(df_bot[indicador.coluna_valor].max()) if len(df_bot) else 0.0
    x_max_bot = max_bot if max_bot > 0 else (overall_max * 0.10 if overall_max > 0 else 1.0)
    ax2.set_xlim(0, x_max_bot * 1.15)
    offset_bot = x_max_bot * 0.02
    for rect, v in zip(bars2, df_bot[indicador.coluna_valor]):
        ax2.text(float(v) + offset_bot, rect.get_y() + rect.get_height() / 2, _fmt_num(float(v), 3),
                 va="center", fontsize=8)

    fig.suptitle(f"{indicador.titulo} (municípios)", fontsize=16, y=1.02)
    # tight_layout pode reclamar quando há legenda/axes extras; manter para melhor encaixe.
    fig.tight_layout()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def gerar_mapa(
    gdf: gpd.GeoDataFrame,
    indicador: Indicador,
    out_path: Path,
) -> None:
    gdf_plot = gdf[["geometry", "NM_MUN", indicador.coluna_valor]].dropna().copy()

    fig, ax = plt.subplots(1, 1, figsize=(11.7, 8.3))  # ~A4 landscape
    gdf_plot.plot(
        column=indicador.coluna_valor,
        cmap=indicador.cmap,
        legend=True,
        ax=ax,
        linewidth=0.15,
        edgecolor="#666666",
    )

    ax.set_title(f"{indicador.titulo} — escala municipal", fontsize=14)
    ax.axis("off")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def indicadores_estaduais(gdf: gpd.GeoDataFrame) -> dict[str, float]:
    area_total = float(gdf["Area_km2"].sum())
    pop_total = float(gdf["Populacao"].sum())
    ext_total = float(gdf["Ext_Total"].sum())

    dens_area = ext_total / area_total if area_total else float("nan")
    dens_pop = ext_total / (pop_total / 1000.0) if pop_total else float("nan")

    return {
        "area_total": area_total,
        "pop_total": pop_total,
        "ext_total": ext_total,
        "dens_area": dens_area,
        "dens_pop": dens_pop,
    }


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Cm(0.15)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _top_bottom_list_text(
    gdf: gpd.GeoDataFrame,
    indicador: Indicador,
    n: int = 10,
    label_col: str = "NM_MUN",
) -> tuple[str, str]:
    df = gdf[[label_col, indicador.coluna_valor]].copy()
    df[indicador.coluna_valor] = pd.to_numeric(df[indicador.coluna_valor], errors="coerce")
    df = df.dropna(subset=[indicador.coluna_valor])

    top = df.nlargest(n, indicador.coluna_valor)
    bot = df.nsmallest(n, indicador.coluna_valor)

    # Definir precisão por indicador
    if indicador.coluna_valor == "Ext_Total":
        dec = 1
    elif indicador.coluna_valor == "Dens_Area_Total":
        dec = 3
    else:
        dec = 3

    top_items = [
        f"{row[label_col]} ({_fmt_num(float(row[indicador.coluna_valor]), dec)} {indicador.unidade})"
        for _, row in top.iterrows()
    ]
    bot_items = [
        f"{row[label_col]} ({_fmt_num(float(row[indicador.coluna_valor]), dec)} {indicador.unidade})"
        for _, row in bot.iterrows()
    ]

    return "; ".join(top_items), "; ".join(bot_items)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"[Figura não encontrada: {image_path.name}]")
        return

    # Largura padrão para caber na página
    doc.add_picture(str(image_path), width=Cm(16.5))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def escrever_docx(gdf: gpd.GeoDataFrame, figuras: dict[str, dict[str, Path]]) -> None:
    doc = Document()

    # Capa simples
    title = doc.add_heading("Nota Técnica — Indicadores de Densidade de Malha Rodoviária", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subt = doc.add_paragraph(
        "Estado de São Paulo | Escala municipal, Regiões Geográficas Imediatas e Intermediárias (IBGE)"
    )
    subt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    dt = doc.add_paragraph(f"Data: {date.today().strftime('%d/%m/%Y')}")
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 1. Introdução
    add_heading(doc, "1. Introdução", 1)
    add_paragraph(
        doc,
        "O Estado de São Paulo concentra a maior dinâmica econômica do país e desempenha papel central na logística nacional, "
        "com elevada circulação de pessoas e cargas e forte integração entre áreas metropolitanas, polos regionais e fronteiras interestaduais. "
        "Nesse contexto, a infraestrutura rodoviária é determinante para competitividade, segurança viária, acesso a serviços e coesão territorial.",
    )
    add_paragraph(
        doc,
        "Esta nota técnica apresenta indicadores de densidade de malha rodoviária para apoiar o diagnóstico territorial e a priorização de "
        "diretrizes, programas e projetos no Plano de Logística e Investimentos (PLI). Os resultados são apresentados prioritariamente na escala municipal "
        "— recorte mais aderente à identificação de disparidades intrarregionais e à estruturação de carteiras de intervenção — e, de forma complementar, "
        "por Regiões Geográficas Imediatas e Intermediárias (IBGE).",
    )
    add_paragraph(
        doc,
        "O foco deste material é responder, de forma objetiva e territorializada, a perguntas como: onde a infraestrutura está mais concentrada; "
        "onde a cobertura territorial é relativamente baixa; e onde a disponibilidade por habitante é menor. Esses indicadores funcionam como ferramenta "
        "de triagem e priorização: apontam onde vale aprofundar estudos, organizar programas e refinar critérios de seleção de projetos no PLI.",
    )

    # 2. Objetivos
    add_heading(doc, "2. Objetivos", 1)
    add_paragraph(
        doc,
        "Para responder às perguntas apresentadas na introdução, foram delineados os seguintes objetivos: (i) caracterizar a distribuição espacial da "
        "malha rodoviária no território paulista; (ii) quantificar a extensão da infraestrutura e sua densidade em relação à área e à população; (iii) "
        "comparar padrões territoriais associados a concessões, hierarquia funcional e contexto urbano-rural; e (iv) evidenciar disparidades municipais e "
        "regionais que orientem critérios territoriais de priorização e a seleção de temas e projetos no PLI.",
    )

    # 3. Dados e metodologia (breve)
    add_heading(doc, "3. Dados e metodologia (síntese)", 1)
    add_paragraph(
        doc,
        "Foram integradas bases vetoriais e estatísticas de fontes oficiais. A malha rodoviária foi tratada a partir de base vetorial do "
        "DER/SP, com agregação da extensão de trechos por unidade territorial. As unidades de análise territorial (municípios, Regiões Geográficas "
        "Imediatas e Intermediárias) foram obtidas do IBGE; a população municipal foi incorporada a partir do Censo 2022."
    )
    add_paragraph(
        doc,
        "Para apoiar leituras urbano-rurais, a classificação de trechos considera a proporção de cada segmento em áreas urbanizadas do IBGE, "
        "com categorias urbano, misto e rural. Os indicadores centrais são calculados como: (i) extensão total; (ii) densidade por área (km/km²); "
        "e (iii) densidade por população (km/10.000 habitantes)."
    )
    add_paragraph(doc, "O detalhamento metodológico (regras de classificação, projeções e controles de qualidade) encontra-se no Anexo A.")

    # 4. Resultados
    add_heading(doc, "4. Resultados", 1)

    # 4.1 Visão geral do estado
    add_heading(doc, "4.1 Visão geral (Estado de São Paulo)", 2)
    est = indicadores_estaduais(gdf)
    add_paragraph(
        doc,
        "Na visão agregada do estado (soma das unidades municipais), obtêm-se os seguintes parâmetros de referência para diagnóstico e monitoramento: "
        f"área territorial aproximada de {_fmt_num(est['area_total'], 0)} km², população de {_fmt_int(est['pop_total'])} habitantes "
        f"(Censo 2022) e extensão total de rodovias de {_fmt_num(est['ext_total'], 0)} km. A partir desses totais, a densidade estadual "
        f"por área é de {_fmt_num(est['dens_area'], 4)} km/10000km² e a densidade por população é de {_fmt_num(est['dens_pop'], 3)} km/10.000hab."
    )
    add_paragraph(
        doc,
        "Esses valores médios estaduais servem como referência de comparação. A análise municipal a seguir permite identificar onde a infraestrutura "
        "está concentrada, onde é relativamente escassa frente ao território e onde há maior pressão populacional sobre a rede, fornecendo evidência "
        "para critérios territoriais no PLI.",
    )

    # 4.2 Indicadores
    add_heading(doc, "4.2 Indicadores (escala municipal)", 2)

    for idx, indicador in enumerate(INDICADORES, 1):
        add_heading(doc, f"4.2.{idx} {indicador.titulo}", 3)
        add_paragraph(
            doc,
            "Definição e interpretação. " + indicador.descricao,
        )

        add_paragraph(
            doc,
            "Aplicação no PLI. O indicador apoia: (i) diagnóstico e comunicação de disparidades territoriais; (ii) priorização de aprofundamentos "
            "(demanda, capacidade, segurança viária, logística, conectividade e manutenção); e (iii) definição de indicadores de acompanhamento. "
            "No PLI, recomenda-se seu uso no capítulo de caracterização do modal rodoviário (diagnóstico), na formulação de diretrizes e programas "
            "(critérios territoriais de priorização) e no monitoramento de resultados da carteira de projetos.",
        )

        add_paragraph(doc, "Equação/definição operacional do indicador:")
        eq = doc.add_paragraph(indicador.equacao)
        if eq.runs:
            eq.runs[0].italic = True

        top_txt, bot_txt = _top_bottom_list_text(gdf, indicador, n=10)
        add_paragraph(doc, "Top 10 municípios (maiores valores): " + top_txt + ".")
        add_paragraph(doc, "Bottom 10 municípios (menores valores): " + bot_txt + ".")

        # Comentário técnico por indicador
        if indicador.coluna_valor == "Ext_Total":
            add_paragraph(
                doc,
                "Comentário técnico. A extensão total é um indicador de porte absoluto de rede: municípios com maior malha tendem a concentrar eixos estruturantes "
                "e maior conectividade territorial. Já os menores valores — inclusive valores nulos — indicam ausência de trechos na base considerada para aqueles "
                "municípios (o que pode refletir municípios de pequena área e/ou limitações de cobertura para vias municipais). Para o PLI, o achado reforça a necessidade "
                "de tratar conjuntamente a rede estadual e as conexões municipais/locais, especialmente quando o objetivo é garantir acessibilidade territorial e integração "
                "de cadeias logísticas."
            )
            add_paragraph(
                doc,
                "Direcionamentos para o PLI. Recomenda-se: fortalecer programas de melhoria e manutenção de ligações locais que alimentam eixos regionais; priorizar "
                "intervenções de conectividade (acessos, travessias e ligações a polos de serviços e equipamentos logísticos); e qualificar a base de dados municipal "
                "(ex.: integração com Rotas Rurais) para reduzir incertezas de diagnóstico em áreas rurais."
            )
        elif indicador.coluna_valor == "Dens_Area_Total":
            add_paragraph(
                doc,
                "Comentário técnico. A densidade por área traduz capilaridade territorial: valores elevados sugerem maior cobertura por km², enquanto valores baixos podem "
                "sinalizar menor cobertura territorial (potenciais lacunas de acessibilidade), condicionadas por relevo, uso do solo e padrão de urbanização. A comparação "
                "Top/Bottom evidencia o contraste entre municípios pequenos e densamente servidos e municípios extensos com rede mais rarefeita."
            )
            add_paragraph(
                doc,
                "Direcionamentos para o PLI. O indicador pode orientar critérios territoriais de priorização: selecionar áreas com densidade baixa para aprofundar análises de "
                "conectividade e acesso a serviços; orientar investimentos em ligações estratégicas (corredores de integração regional, rotas de escoamento e acessos a "
                "equipamentos logísticos) e estruturar metas/indicadores de cobertura territorial em recortes específicos, quando aplicável."
            )
        else:
            add_paragraph(
                doc,
                "Comentário técnico. A densidade por população (km/10.000hab) não mede demanda ou nível de serviço; indica disponibilidade relativa de infraestrutura frente ao "
                "tamanho da população. Valores muito altos tendem a ocorrer em municípios com baixa população, mesmo que a rede seja pequena; valores baixos sugerem maior pressão "
                "populacional sobre a rede e potencial necessidade de gestão de capacidade, segurança viária e alternativas modais."
            )
            add_paragraph(
                doc,
                "Direcionamentos para o PLI. Em municípios com baixa disponibilidade por habitante, recomenda-se priorizar ações de capacidade e segurança em corredores críticos "
                "(duplicações, faixas adicionais, eliminação de conflitos, travessias urbanas), além de medidas de gestão de demanda e logística urbana/regional. Em municípios com "
                "valores muito altos, recomenda-se evitar interpretação automática de “excesso de infraestrutura”: deve-se avaliar conectividade, qualidade da via, função logística e "
                "custos de manutenção, para orientar intervenções mais eficientes."
            )

        chart_path = figuras[indicador.key]["grafico"]
        map_path = figuras[indicador.key]["mapa"]

        add_figure(doc, chart_path, f"Figura {idx}. Top 10 e Bottom 10 municípios — {indicador.titulo}.")
        add_figure(doc, map_path, f"Figura {idx + 10}. Mapa coroplético municipal — {indicador.titulo}.")

    # Conclusão
    add_heading(doc, "5. Conclusão", 1)
    add_paragraph(
        doc,
        "Os indicadores de extensão e densidade permitem leitura integrada do sistema rodoviário em três perspectivas complementares: porte absoluto de infraestrutura (km), "
        "cobertura territorial (km/10000km²) e disponibilidade relativa frente à população (km/10.000hab). A leitura conjunta dos Top/Bottom 10 evidencia heterogeneidade municipal, "
        "o que reforça a necessidade de critérios territoriais explícitos no PLI para orientar priorização, aprofundamento de estudos e monitoramento de resultados da carteira "
        "de programas e projetos."
    )
    add_paragraph(
        doc,
        "Como encaminhamento, recomenda-se utilizar estes indicadores como camada de diagnóstico e triagem: municípios com baixa densidade por área podem ser priorizados para análises "
        "de conectividade e acesso a serviços; municípios com baixa disponibilidade por habitante podem demandar ações de capacidade, segurança e gestão de demanda; e municípios com valores "
        "nulos/baixos devem ser verificados quanto à completude da base municipal, evitando conclusões derivadas de lacunas de dados."
    )

    # Anexo
    doc.add_page_break()
    add_heading(doc, "Anexo A — Metodologia detalhada (resumo expandido)", 1)
    add_paragraph(
        doc,
        "(i) A base vetorial da malha rodoviária do DER/SP foi padronizada e validada quanto a campos essenciais, com tipificação por administração (concessionada/não concessionada) "
        "e por hierarquia (SP/SPA/SPI/BR delegada). (ii) Para a leitura urbano-rural, aplicou-se classificação baseada em Áreas Urbanizadas do IBGE: para cada trecho, calculou-se a "
        "proporção do seu comprimento sobre a mancha urbanizada e classificou-se como urbano (>50%), misto (10–50%) ou rural (<10%). (iii) A agregação municipal e regional foi feita por soma "
        "de extensão (km), produzindo também componentes por concessão, tipo e localização. (iv) Para área (km²) por unidade territorial, utilizou-se o atributo do IBGE quando disponível; caso contrário, "
        "calculou-se por reprojeção para EPSG:5880. (v) As densidades por área e por população foram calculadas conforme as equações apresentadas no corpo do relatório."
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    municipios = carregar_municipios()

    figuras: dict[str, dict[str, Path]] = {}
    for indicador in INDICADORES:
        chart_path = FIG_DIR / f"{indicador.key}_top_bottom.png"
        map_path = FIG_DIR / f"{indicador.key}_mapa.png"

        # Reutilizar figuras já geradas (evita retrabalho quando só o texto muda)
        if not chart_path.exists():
            gerar_grafico_top_bottom(municipios, indicador, chart_path)
        if not map_path.exists():
            gerar_mapa(municipios, indicador, map_path)

        figuras[indicador.key] = {"grafico": chart_path, "mapa": map_path}

    escrever_docx(municipios, figuras)

    print("OK")
    print(f"DOCX gerado: {OUT_DOCX}")
    print(f"Figuras: {FIG_DIR}")


if __name__ == "__main__":
    main()
