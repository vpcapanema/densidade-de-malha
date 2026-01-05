#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def adicionar_header_oficio(doc):
    """Adiciona cabeçalho padrão ao documento"""
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("[Brasão e Cabeçalho da SEMIL]")
    run.font.size = Pt(11)
    run.font.italic = True

def criar_oficio_formal():
    """Cria Ofício Versão 1 - Formal/Protocolar"""
    doc = Document()
    
    adicionar_header_oficio(doc)
    
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("OFÍCIO n.º [___]/2026/SEMIL")
    run.font.bold = True
    run.font.size = Pt(12)
    
    # Data e local
    data = doc.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_run = data.add_run("São Paulo, [data]")
    data_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espaço
    
    # Saudação
    doc.add_paragraph("Senhor Secretário,")
    
    # Corpo
    corpo1 = doc.add_paragraph(
        "Vimos respeitosamente solicitar apoio técnico-administrativo da Secretaria de Estado "
        "de Agricultura e Abastecimento em relação ao programa \"Rotas Rurais\", conforme descrito a seguir."
    )
    corpo1.paragraph_format.line_spacing = 1.5
    
    corpo2 = doc.add_paragraph(
        "A Secretaria de Meio Ambiente, Infraestrutura e Logística (SEMIL) encontra-se envolvida "
        "na elaboração do Plano de Logística e Investimentos (PLI-SP 2050), cujo Produto 3 aborda a "
        "caracterização da infraestrutura rodoviária estadual. Nesta fase de desenvolvimento do documento, "
        "foi identificado pelo Departamento de Estradas de Rodagem (DER/SP) que a caracterização completa "
        "da malha rodoviária municipal exige consulta a fontes de dados adicionais."
    )
    corpo2.paragraph_format.line_spacing = 1.5
    
    corpo3 = doc.add_paragraph(
        "Nesse sentido, tomamos conhecimento de que a Secretaria de Agricultura e Abastecimento "
        "mantém o programa \"Rotas Rurais\", disponível em "
        "https://arcgisrurais.agricultura.sp.gov.br/portal/apps/sites/#/iea-rotas-rurais, "
        "que incorpora informações significativamente mais abrangentes sobre a rede rodoviária "
        "municipal estadual."
    )
    corpo3.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    # Solicitações
    solicitacoes = doc.add_paragraph("Solicitamos cordialmente:")
    solicitacoes.runs[0].font.bold = True
    solicitacoes.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph(
        "Confirmar a existência de base de dados vetoriais consolidada referente às rotas rurais, "
        "com especificações técnicas (formato, projeção cartográfica, data de atualização e cobertura geográfica);",
        style='List Number'
    )
    
    doc.add_paragraph(
        "Conhecer as possibilidades de compartilhamento dessa base de dados com a SEMIL, "
        "para fins de incorporação ao Produto 3 do PLI-SP 2050;",
        style='List Number'
    )
    
    doc.add_paragraph(
        "Indicar o setor/unidade competente responsável pela gestão e disponibilização dessa informação.",
        style='List Number'
    )
    
    doc.add_paragraph()  # Espaço
    
    # Fechamento
    fechamento = doc.add_paragraph(
        "Colocamo-nos à disposição para esclarecer dúvidas adicionais ou formalizar acordo de "
        "compartilhamento de dados que se mostre necessário."
    )
    fechamento.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    # Assinatura
    doc.add_paragraph("Atenciosamente,")
    doc.add_paragraph()
    doc.add_paragraph("[Nome e Cargo]")
    doc.add_paragraph("Secretaria de Meio Ambiente, Infraestrutura e Logística (SEMIL)")
    
    return doc

def criar_oficio_tecnico():
    """Cria Ofício Versão 2 - Técnica/Detalhada"""
    doc = Document()
    
    adicionar_header_oficio(doc)
    
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("OFÍCIO n.º [___]/2026/SEMIL - DIRETORIA DE LOGÍSTICA")
    run.font.bold = True
    run.font.size = Pt(12)
    
    # Data
    data = doc.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_run = data.add_run("São Paulo, [data]")
    data_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espaço
    
    # Assunto
    assunto = doc.add_paragraph()
    assunto.add_run("Assunto: ").bold = True
    assunto.add_run("Solicitação de Dados Vetoriais - Programa Rotas Rurais")
    
    doc.add_paragraph()  # Espaço
    
    # Saudação
    doc.add_paragraph("Senhor Secretário,")
    
    # Corpo
    corpo1 = doc.add_paragraph(
        "No contexto da elaboração do Plano de Logística e Investimentos (PLI-SP 2050), "
        "a Secretaria de Meio Ambiente, Infraestrutura e Logística (SEMIL), em articulação com o "
        "Departamento de Estradas de Rodagem (DER/SP), identificou necessidade de aprofundar a "
        "análise da malha rodoviária municipal estadual, especialmente com vistas à caracterização "
        "de déficits de infraestrutura em áreas rurais e análise de densidade de malha por diferentes "
        "unidades administrativas."
    )
    corpo1.paragraph_format.line_spacing = 1.5
    
    corpo2 = doc.add_paragraph(
        "A análise técnica conduzida pelo DER/SP indicou que as bases de dados convencionalmente "
        "utilizadas (OpenStreetMap, IBGE) apresentam lacunas significativas em relação à cobertura e "
        "precisão da malha rodoviária municipal, particularmente em áreas rurais e de menor densidade demográfica."
    )
    corpo2.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    # Solicitações
    solicitacoes = doc.add_paragraph("Neste sentido, solicitamos informações técnicas sobre:")
    solicitacoes.runs[0].font.bold = True
    solicitacoes.paragraph_format.line_spacing = 1.5
    
    # Item 1
    item1 = doc.add_paragraph()
    item1.add_run("Disponibilidade de dados: ").bold = True
    item1.add_run(
        "Confirmar se existe base de dados vetoriais consolidada referente ao programa "
        "\"Rotas Rurais\", com cobertura estadual completa;"
    )
    item1.paragraph_format.left_indent = Inches(0.5)
    
    # Item 2
    item2 = doc.add_paragraph()
    item2.add_run("Especificações técnicas: ").bold = True
    item2.add_run("Providenciar informações sobre:")
    item2.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph("Formato de disponibilização (shapefile, GeoJSON, geodatabase, etc.)", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    doc.add_paragraph("Sistema de projeção cartográfica adotado", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    doc.add_paragraph("Data da última atualização e periodicidade de revisão", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    doc.add_paragraph("Escala ou precisão posicional dos dados", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    doc.add_paragraph("Cobertura geográfica (municípios/regiões atendidas)", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    
    # Item 3
    item3 = doc.add_paragraph()
    item3.add_run("Condições de acesso: ").bold = True
    item3.add_run("Esclarecer as possibilidades de:")
    item3.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph("Acesso direto aos dados em formato vetorial editável", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    doc.add_paragraph("Licença ou termo de uso aplicável", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    doc.add_paragraph("Integrações técnicas ou APIs disponíveis", 
                     style='List Bullet').paragraph_format.left_indent = Inches(1.0)
    
    # Item 4
    item4 = doc.add_paragraph()
    item4.add_run("Responsabilidade técnica: ").bold = True
    item4.add_run(
        "Indicar setor/unidade responsável pela curadoria e manutenção dos dados, "
        "para fins de contato e possíveis ajustes/atualizações futuras."
    )
    item4.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()  # Espaço
    
    # Fechamento
    fechamento = doc.add_paragraph(
        "Ressaltamos que a incorporação dessas informações contribuirá significativamente para a "
        "qualidade técnica do documento e para políticas públicas de investimento em infraestrutura "
        "rodoviária municipal."
    )
    fechamento.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    fechamento2 = doc.add_paragraph(
        "Colocamo-nos à disposição para discussões técnicas, execução de termos de cooperação ou "
        "qualquer procedimento administrativo necessário."
    )
    fechamento2.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    # Assinatura
    doc.add_paragraph("Atenciosamente,")
    doc.add_paragraph()
    doc.add_paragraph("[Nome e Cargo]")
    doc.add_paragraph("Secretaria de Meio Ambiente, Infraestrutura e Logística (SEMIL)")
    doc.add_paragraph("[E-mail e Telefone para contato técnico]")
    
    return doc

def criar_oficio_colaborativo():
    """Cria Ofício Versão 3 - Colaborativa/Amigável"""
    doc = Document()
    
    adicionar_header_oficio(doc)
    
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("OFÍCIO n.º [___]/2026/SEMIL")
    run.font.bold = True
    run.font.size = Pt(12)
    
    # Data
    data = doc.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_run = data.add_run("São Paulo, [data]")
    data_run.font.size = Pt(11)
    
    doc.add_paragraph()  # Espaço
    
    # Saudação
    doc.add_paragraph("Prezado Secretário,")
    
    # Corpo
    corpo1 = doc.add_paragraph(
        "Esperamos que se encontre bem! Vimos até você com uma solicitação que acreditamos poder "
        "beneficiar ambas as secretarias e, sobretudo, a qualidade das políticas públicas para o "
        "desenvolvimento rural paulista."
    )
    corpo1.paragraph_format.line_spacing = 1.5
    
    corpo2 = doc.add_paragraph(
        "Estamos finalizando o Plano de Logística e Investimentos (PLI-SP 2050), um documento que "
        "busca mapear e planejar a infraestrutura de transporte estadual para as próximas décadas. "
        "Em uma das análises que estamos desenvolvendo - justamente sobre a rede rodoviária municipal - "
        "descobrimos que vocês têm um trabalho fantástico em andamento: o programa Rotas Rurais."
    )
    corpo2.paragraph_format.line_spacing = 1.5
    
    corpo3 = doc.add_paragraph(
        "Tomamos conhecimento dessa iniciativa através do DER/SP e ficamos impressionados com a "
        "abrangência dos dados que vocês conseguiram consolidar. Seria de grande valor para nosso "
        "documento conseguir acessar essa base de dados, tanto para enriquecer nossa análise quanto "
        "para reconhecer o esforço da Secretaria de Agricultura em caracterizar a infraestrutura que "
        "sustenta as atividades agrícolas e rurais do Estado."
    )
    corpo3.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    # Perguntas
    perguntas = doc.add_paragraph("Nossa pergunta é simples:")
    perguntas.runs[0].font.bold = True
    perguntas.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph(
        "Vocês disponibilizariam a base de dados vetorial consolidada do programa Rotas Rurais?",
        style='List Number'
    )
    
    doc.add_paragraph(
        "Qual seria o melhor caminho (formal ou técnico) para que conseguíssemos acessar e utilizar esses dados?",
        style='List Number'
    )
    
    doc.add_paragraph()  # Espaço
    
    # Fechamento
    fechamento = doc.add_paragraph(
        "Entendemos que isso pode gerar uma sinergia interessante: dados melhores para o nosso documento, "
        "visibilidade para o programa Rotas Rurais, e subsídios mais sólidos para decisões sobre investimentos "
        "em infraestrutura rural."
    )
    fechamento.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    fechamento2 = doc.add_paragraph(
        "Se for possível, gostaríamos também de conhecer a equipe técnica responsável, pois talvez haja "
        "oportunidades de colaboração contínua entre nossas secretarias."
    )
    fechamento2.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    fechamento3 = doc.add_paragraph(
        "Fico no aguardo de um retorno. Estarei feliz em discutir detalhes técnicos, assinar termos de "
        "compartilhamento ou qualquer outra formalidade que se faça necessária."
    )
    fechamento3.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()  # Espaço
    
    # Assinatura
    doc.add_paragraph("Cordialmente,")
    doc.add_paragraph()
    doc.add_paragraph("[Nome e Cargo]")
    doc.add_paragraph("Secretaria de Meio Ambiente, Infraestrutura e Logística (SEMIL)")
    doc.add_paragraph("[E-mail] / [Telefone]")
    
    return doc

if __name__ == "__main__":
    # Criar documentos
    print("Gerando Ofício Versão 1 - Formal/Protocolar...")
    doc1 = criar_oficio_formal()
    doc1.save("D:\\densidade _de_malha\\Oficio_v1_Formal.docx")
    print("✓ Oficio_v1_Formal.docx criado com sucesso!")
    
    print("\nGerando Ofício Versão 2 - Técnica/Detalhada...")
    doc2 = criar_oficio_tecnico()
    doc2.save("D:\\densidade _de_malha\\Oficio_v2_Tecnica.docx")
    print("✓ Oficio_v2_Tecnica.docx criado com sucesso!")
    
    print("\nGerando Ofício Versão 3 - Colaborativa/Amigável...")
    doc3 = criar_oficio_colaborativo()
    doc3.save("D:\\densidade _de_malha\\Oficio_v3_Colaborativa.docx")
    print("✓ Oficio_v3_Colaborativa.docx criado com sucesso!")
    
    print("\n" + "="*50)
    print("Todos os ofícios foram exportados para Word!")
    print("="*50)
